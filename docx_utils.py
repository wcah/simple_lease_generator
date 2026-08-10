import io
import re

from docx import Document as Doc
from docx.document import Document

PLACEHOLDER_PATTERN = re.compile(r"\[\[([A-Za-z][A-Za-z0-9_]*)\]\]")


def read_docx(source: str | io.BytesIO) -> Document:
    return Doc(source)


def get_text(doc: Document) -> str:
    return "\n".join(para.text for para in doc.paragraphs)


def find_placeholders(doc: Document) -> list[str]:
    found = set()
    for para in doc.paragraphs:
        found.update(match.upper() for match in PLACEHOLDER_PATTERN.findall(para.text))
    return sorted(found)


def write_docx(doc: Document, target: str | io.BytesIO, source_path: str | None = None) -> None:
    if source_path and isinstance(target, str) and target == source_path:
        raise ValueError(f"Refusing to overwrite source file: {source_path}")
    doc.save(target)


def replace_placeholder(doc: Document, placeholder: str, replacement: str) -> None:
    token = f"[[{placeholder.upper()}]]"
    for para in doc.paragraphs:
        if token in para.text:
            for run in para.runs:
                if token in run.text:
                    run.text = run.text.replace(token, replacement)



